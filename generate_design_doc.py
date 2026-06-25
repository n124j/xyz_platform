"""
Generate XYZ Platform Design Document as a Microsoft Word (.docx) file
with flowcharts and data flow diagrams using matplotlib.

Usage:
    python generate_design_doc.py
"""
import os
import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DOCX = os.path.join(BASE_DIR, "XYZ_Platform_Design_Document.docx")
IMG_DIR = os.path.join(BASE_DIR, "doc_images")
os.makedirs(IMG_DIR, exist_ok=True)

COLORS = {
    "blue": "#2563EB",
    "dark_blue": "#1E3A5F",
    "light_blue": "#DBEAFE",
    "green": "#059669",
    "light_green": "#D1FAE5",
    "orange": "#D97706",
    "light_orange": "#FEF3C7",
    "red": "#DC2626",
    "light_red": "#FEE2E2",
    "gray": "#6B7280",
    "light_gray": "#F3F4F6",
    "purple": "#7C3AED",
    "light_purple": "#EDE9FE",
    "white": "#FFFFFF",
    "black": "#111827",
}


def draw_rounded_box(ax, x, y, w, h, text, facecolor, edgecolor, fontsize=8, fontweight="normal", textcolor="black"):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.1",
        facecolor=facecolor,
        edgecolor=edgecolor,
        linewidth=1.5,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center",
        fontsize=fontsize, fontweight=fontweight,
        color=textcolor, wrap=True,
    )


def draw_arrow(ax, x1, y1, x2, y2, color="#374151"):
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(arrowstyle="->", color=color, lw=1.5),
    )


def generate_system_architecture_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("System Architecture Diagram", fontsize=14, fontweight="bold", pad=20)

    draw_rounded_box(ax, 4, 7, 4, 0.7, "Browser / Staff UI", COLORS["light_blue"], COLORS["blue"], fontsize=10, fontweight="bold")

    draw_rounded_box(ax, 3.5, 5.8, 5, 0.7, "Nginx (TLS + Rate Limiting)", COLORS["light_orange"], COLORS["orange"], fontsize=9, fontweight="bold")

    draw_rounded_box(ax, 0.5, 4, 2.5, 0.7, "Django Views\n(Templates)", COLORS["light_green"], COLORS["green"], fontsize=8)
    draw_rounded_box(ax, 3.5, 4, 2.5, 0.7, "Plotly Dash\n(Embedded)", COLORS["light_purple"], COLORS["purple"], fontsize=8)
    draw_rounded_box(ax, 6.5, 4, 2.5, 0.7, "DRF API\n(/api/v1/)", COLORS["light_blue"], COLORS["blue"], fontsize=8)
    draw_rounded_box(ax, 9.5, 4, 2, 0.7, "Django\nChannels", COLORS["light_orange"], COLORS["orange"], fontsize=8)

    draw_rounded_box(ax, 3, 2.5, 6, 0.7, "PostgreSQL 16 (accounts, portfolio, analytics, etl_monitor)", COLORS["light_gray"], COLORS["gray"], fontsize=8, fontweight="bold")

    draw_rounded_box(ax, 0.5, 0.8, 2.2, 0.7, "Celery Worker\n(Tasks)", COLORS["light_green"], COLORS["green"], fontsize=7)
    draw_rounded_box(ax, 3.2, 0.8, 2.2, 0.7, "Celery Beat\n(Scheduler)", COLORS["light_green"], COLORS["green"], fontsize=7)
    draw_rounded_box(ax, 5.9, 0.8, 2.5, 0.7, "Airflow\n(DAGs)", COLORS["light_orange"], COLORS["orange"], fontsize=7)
    draw_rounded_box(ax, 8.9, 0.8, 2.5, 0.7, "Redis\n(Cache + Broker)", COLORS["light_red"], COLORS["red"], fontsize=7)

    draw_arrow(ax, 6, 7, 6, 6.55)
    draw_arrow(ax, 6, 5.8, 1.75, 4.75)
    draw_arrow(ax, 6, 5.8, 4.75, 4.75)
    draw_arrow(ax, 6, 5.8, 7.75, 4.75)
    draw_arrow(ax, 6, 5.8, 10.5, 4.75)

    draw_arrow(ax, 1.75, 4, 5, 3.25)
    draw_arrow(ax, 4.75, 4, 5.5, 3.25)
    draw_arrow(ax, 7.75, 4, 6.5, 3.25)

    draw_arrow(ax, 4.5, 2.5, 1.6, 1.55)
    draw_arrow(ax, 5, 2.5, 4.3, 1.55)
    draw_arrow(ax, 7, 2.5, 7.15, 1.55)
    draw_arrow(ax, 10.15, 1.55, 10.15, 2.5)

    path = os.path.join(IMG_DIR, "system_architecture.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def generate_data_flow_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(12, 9))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Data Flow Diagram", fontsize=14, fontweight="bold", pad=20)

    draw_rounded_box(ax, 0.3, 7.5, 2.5, 0.8, "Market Data\nProviders", COLORS["light_orange"], COLORS["orange"], fontsize=8, fontweight="bold")
    draw_rounded_box(ax, 4.5, 7.5, 2.5, 0.8, "Custody / OMS\nSystems", COLORS["light_orange"], COLORS["orange"], fontsize=8, fontweight="bold")
    draw_rounded_box(ax, 8.8, 7.5, 2.5, 0.8, "Bloomberg /\nRefinitiv", COLORS["light_orange"], COLORS["orange"], fontsize=8, fontweight="bold")

    draw_rounded_box(ax, 0.3, 5.5, 3.2, 1, "Market Data DAG\n(every 15 min)\nFetch OHLCV + FX + Benchmarks", COLORS["light_blue"], COLORS["blue"], fontsize=7)
    draw_rounded_box(ax, 4.2, 5.5, 3.2, 1, "Portfolio ETL DAG\n(daily 18:30 ET)\nExtract/Transform/Load", COLORS["light_blue"], COLORS["blue"], fontsize=7)
    draw_rounded_box(ax, 8.2, 5.5, 3.2, 1, "Risk Report DAG\n(daily 19:00 ET)\nVaR / Stress Tests", COLORS["light_blue"], COLORS["blue"], fontsize=7)

    draw_rounded_box(ax, 0.3, 3.5, 2, 0.8, "MarketData\nTable", COLORS["light_gray"], COLORS["gray"], fontsize=7, fontweight="bold")
    draw_rounded_box(ax, 2.8, 3.5, 2, 0.8, "Holding &\nAccount", COLORS["light_gray"], COLORS["gray"], fontsize=7, fontweight="bold")
    draw_rounded_box(ax, 5.3, 3.5, 2, 0.8, "Portfolio\nSnapshot", COLORS["light_gray"], COLORS["gray"], fontsize=7, fontweight="bold")
    draw_rounded_box(ax, 7.8, 3.5, 2, 0.8, "RiskMetric\nTable", COLORS["light_gray"], COLORS["gray"], fontsize=7, fontweight="bold")
    draw_rounded_box(ax, 10.3, 3.5, 1.3, 0.8, "Benchmark\nReturn", COLORS["light_gray"], COLORS["gray"], fontsize=7, fontweight="bold")

    draw_rounded_box(ax, 0.3, 1.5, 2.5, 0.8, "Portfolio\nDashboard", COLORS["light_green"], COLORS["green"], fontsize=8, fontweight="bold")
    draw_rounded_box(ax, 3.5, 1.5, 2.5, 0.8, "Risk Analytics\nDashboard", COLORS["light_green"], COLORS["green"], fontsize=8, fontweight="bold")
    draw_rounded_box(ax, 6.7, 1.5, 2.5, 0.8, "Client &\nAccount Views", COLORS["light_green"], COLORS["green"], fontsize=8, fontweight="bold")
    draw_rounded_box(ax, 9.8, 1.5, 1.8, 0.8, "REST\nAPI", COLORS["light_green"], COLORS["green"], fontsize=8, fontweight="bold")

    draw_rounded_box(ax, 4, 0, 4, 0.7, "Browser / Staff Users", COLORS["light_blue"], COLORS["blue"], fontsize=9, fontweight="bold")

    draw_arrow(ax, 1.55, 7.5, 1.55, 6.55)
    draw_arrow(ax, 5.75, 7.5, 5.75, 6.55)
    draw_arrow(ax, 10.05, 7.5, 10.05, 6.55)

    draw_arrow(ax, 1.55, 5.5, 1.3, 4.35)
    draw_arrow(ax, 5.8, 5.5, 3.8, 4.35)
    draw_arrow(ax, 5.8, 5.5, 6.3, 4.35)
    draw_arrow(ax, 9.8, 5.5, 8.8, 4.35)
    draw_arrow(ax, 2.5, 5.5, 10.95, 4.35)

    draw_arrow(ax, 7.3, 5.5, 5.8, 6.55, color=COLORS["green"])

    draw_arrow(ax, 1.3, 3.5, 1.55, 2.35)
    draw_arrow(ax, 6.3, 3.5, 1.55, 2.35)
    draw_arrow(ax, 8.8, 3.5, 4.75, 2.35)
    draw_arrow(ax, 3.8, 3.5, 7.95, 2.35)
    draw_arrow(ax, 1.3, 3.5, 10.7, 2.35)

    draw_arrow(ax, 1.55, 1.5, 5, 0.75)
    draw_arrow(ax, 4.75, 1.5, 5.5, 0.75)
    draw_arrow(ax, 7.95, 1.5, 6.5, 0.75)
    draw_arrow(ax, 10.7, 1.5, 7, 0.75)

    ax.text(6, 8.7, "EXTERNAL SOURCES", fontsize=9, fontweight="bold", ha="center", color=COLORS["orange"])
    ax.text(6, 6.9, "AIRFLOW ETL PIPELINES", fontsize=9, fontweight="bold", ha="center", color=COLORS["blue"])
    ax.text(6, 4.7, "POSTGRESQL DATA STORES", fontsize=9, fontweight="bold", ha="center", color=COLORS["gray"])
    ax.text(6, 2.7, "DJANGO PRESENTATION LAYER", fontsize=9, fontweight="bold", ha="center", color=COLORS["green"])

    path = os.path.join(IMG_DIR, "data_flow_diagram.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def generate_etl_flowchart():
    fig, axes = plt.subplots(1, 3, figsize=(14, 8))
    fig.suptitle("ETL Pipeline Flowcharts", fontsize=14, fontweight="bold")

    for ax_item in axes:
        ax_item.set_xlim(0, 4)
        ax_item.set_ylim(0, 10)
        ax_item.axis("off")

    ax = axes[0]
    ax.set_title("Portfolio ETL DAG\n(Daily 18:30 ET)", fontsize=9, fontweight="bold")
    steps = [
        ("Start", COLORS["light_gray"], COLORS["gray"]),
        ("Extract\nHoldings", COLORS["light_blue"], COLORS["blue"]),
        ("Validate\nHoldings", COLORS["light_orange"], COLORS["orange"]),
        ("Transform\nHoldings", COLORS["light_blue"], COLORS["blue"]),
        ("Load\nHoldings", COLORS["light_blue"], COLORS["blue"]),
        ("Compute\nPortfolio Stats", COLORS["light_purple"], COLORS["purple"]),
        ("Refresh Risk\nMetrics", COLORS["light_purple"], COLORS["purple"]),
        ("Notify\nSuccess", COLORS["light_green"], COLORS["green"]),
    ]
    for i, (label, fc, ec) in enumerate(steps):
        y = 9 - i * 1.1
        draw_rounded_box(ax, 0.8, y, 2.4, 0.7, label, fc, ec, fontsize=7)
        if i > 0:
            draw_arrow(ax, 2, y + 0.75 + 0.4, 2, y + 0.7)

    ax.text(3.5, 9 - 2 * 1.1 + 0.35, "FAIL", fontsize=7, color=COLORS["red"], fontweight="bold")

    ax = axes[1]
    ax.set_title("Market Data DAG\n(Every 15 min)", fontsize=9, fontweight="bold")
    steps2 = [
        ("Start", COLORS["light_gray"], COLORS["gray"]),
        ("Check Market\nOpen", COLORS["light_orange"], COLORS["orange"]),
        ("Fetch Equity\nPrices", COLORS["light_blue"], COLORS["blue"]),
        ("Fetch FX\nRates", COLORS["light_blue"], COLORS["blue"]),
        ("Fetch Benchmark\nLevels", COLORS["light_blue"], COLORS["blue"]),
        ("Validate\nPrices", COLORS["light_orange"], COLORS["orange"]),
        ("Persist\nMarket Data", COLORS["light_blue"], COLORS["blue"]),
        ("Update\nHolding Prices", COLORS["light_purple"], COLORS["purple"]),
        ("Update Account\nValues", COLORS["light_green"], COLORS["green"]),
    ]
    for i, (label, fc, ec) in enumerate(steps2):
        y = 9.3 - i * 1.0
        draw_rounded_box(ax, 0.8, y, 2.4, 0.65, label, fc, ec, fontsize=7)
        if i > 0:
            draw_arrow(ax, 2, y + 0.65 + 0.35, 2, y + 0.65)

    ax = axes[2]
    ax.set_title("Risk Report DAG\n(Daily 19:00 ET)", fontsize=9, fontweight="bold")
    steps3 = [
        ("Start", COLORS["light_gray"], COLORS["gray"]),
        ("Wait for\nPortfolio ETL", COLORS["light_orange"], COLORS["orange"]),
        ("Historical\nVaR", COLORS["light_blue"], COLORS["blue"]),
        ("Parametric\nVaR", COLORS["light_blue"], COLORS["blue"]),
        ("Stress\nTests", COLORS["light_red"], COLORS["red"]),
        ("BHB\nAttribution", COLORS["light_blue"], COLORS["blue"]),
        ("Build Risk\nReport", COLORS["light_purple"], COLORS["purple"]),
        ("Persist\nMetrics", COLORS["light_purple"], COLORS["purple"]),
        ("Distribute\nReports", COLORS["light_green"], COLORS["green"]),
    ]
    for i, (label, fc, ec) in enumerate(steps3):
        y = 9.3 - i * 1.0
        draw_rounded_box(ax, 0.8, y, 2.4, 0.65, label, fc, ec, fontsize=7)
        if i > 0:
            draw_arrow(ax, 2, y + 0.65 + 0.35, 2, y + 0.65)

    path = os.path.join(IMG_DIR, "etl_flowchart.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def generate_er_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.set_title("Entity Relationship Diagram", fontsize=14, fontweight="bold", pad=20)

    def draw_entity(ax_obj, x, y, w, h, title, fields, color, edge):
        box = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.05",
            facecolor=color,
            edgecolor=edge,
            linewidth=2,
        )
        ax_obj.add_patch(box)
        header_box = FancyBboxPatch(
            (x, y + h - 0.45), w, 0.45,
            boxstyle="round,pad=0.05",
            facecolor=edge,
            edgecolor=edge,
            linewidth=1,
        )
        ax_obj.add_patch(header_box)
        ax_obj.text(x + w / 2, y + h - 0.22, title, ha="center", va="center",
                    fontsize=8, fontweight="bold", color="white")
        for i, field in enumerate(fields):
            ax_obj.text(x + 0.1, y + h - 0.65 - i * 0.22, field, fontsize=6, color=COLORS["black"])

    draw_entity(ax, 0.2, 7, 2.8, 2.5, "User (auth_user)",
                ["id (PK)", "username", "email", "is_staff", "is_superuser"],
                COLORS["light_gray"], COLORS["gray"])

    draw_entity(ax, 4, 6.5, 3, 3, "Client",
                ["id (PK)", "client_id (UK)", "name", "email (UK)", "tier", "risk_profile",
                 "relationship_manager (FK)", "kyc_verified", "onboarded_date", "is_active"],
                COLORS["light_blue"], COLORS["blue"])

    draw_entity(ax, 8.5, 6.5, 3, 2.5, "Account",
                ["id (PK)", "account_number (UK)", "client (FK)", "account_type",
                 "market_value", "cash_balance", "ytd_return", "is_active"],
                COLORS["light_blue"], COLORS["blue"])

    draw_entity(ax, 0.2, 3.5, 2.8, 2.5, "Holding",
                ["id (PK)", "account (FK)", "ticker", "asset_class",
                 "quantity", "cost_basis", "market_price",
                 "market_value", "unrealized_pnl"],
                COLORS["light_green"], COLORS["green"])

    draw_entity(ax, 3.5, 3.5, 3, 2.5, "Transaction",
                ["id (PK)", "account (FK)", "transaction_type", "ticker",
                 "trade_date", "settlement_date",
                 "gross_amount", "fees", "net_amount", "reference_number (UK)"],
                COLORS["light_green"], COLORS["green"])

    draw_entity(ax, 7, 3.5, 2.8, 2, "PortfolioSnapshot",
                ["id (PK)", "snapshot_date (UK)", "total_aum",
                 "equity_value", "fi_value", "daily_pnl", "ytd_return"],
                COLORS["light_purple"], COLORS["purple"])

    draw_entity(ax, 10.3, 3.5, 3.2, 2.3, "RiskMetric",
                ["id (PK)", "scope", "reference_id",
                 "calculation_date", "var_95_1d", "sharpe_ratio",
                 "max_drawdown", "beta", "alpha"],
                COLORS["light_orange"], COLORS["orange"])

    draw_entity(ax, 0.2, 0.5, 2.8, 2.2, "MarketData",
                ["id (PK)", "ticker", "price_date",
                 "open/high/low/close", "volume",
                 "UK: (ticker, price_date)"],
                COLORS["light_orange"], COLORS["orange"])

    draw_entity(ax, 3.5, 0.5, 3, 2.2, "DAGRun",
                ["id (PK)", "dag_id", "dag_run_id (UK)",
                 "state", "execution_date", "duration_seconds"],
                COLORS["light_red"], COLORS["red"])

    draw_entity(ax, 7, 0.5, 2.8, 2.2, "TaskInstance",
                ["id (PK)", "dag_run (FK)", "task_id",
                 "state", "duration_seconds",
                 "UK: (dag_run, task_id)"],
                COLORS["light_red"], COLORS["red"])

    draw_entity(ax, 10.3, 0.5, 3.2, 2.2, "PipelineAlert",
                ["id (PK)", "dag_run (FK, null)",
                 "dag_id", "severity", "message",
                 "acknowledged", "acknowledged_by (FK)"],
                COLORS["light_red"], COLORS["red"])

    ax.annotate("", xy=(4, 8), xytext=(3, 8),
                arrowprops=dict(arrowstyle="->", color=COLORS["gray"], lw=1.5))
    ax.text(3.3, 8.15, "RM", fontsize=6, color=COLORS["gray"])

    ax.annotate("", xy=(8.5, 8), xytext=(7, 8),
                arrowprops=dict(arrowstyle="->", color=COLORS["blue"], lw=1.5))
    ax.text(7.5, 8.15, "1:N", fontsize=6, color=COLORS["blue"])

    ax.annotate("", xy=(1.6, 6), xytext=(9, 6.5),
                arrowprops=dict(arrowstyle="->", color=COLORS["green"], lw=1.5))
    ax.text(5, 6.35, "1:N", fontsize=6, color=COLORS["green"])

    ax.annotate("", xy=(5, 6), xytext=(9.5, 6.5),
                arrowprops=dict(arrowstyle="->", color=COLORS["green"], lw=1.5))
    ax.text(7.5, 6.1, "1:N", fontsize=6, color=COLORS["green"])

    ax.annotate("", xy=(8.4, 1.6), xytext=(6.5, 1.6),
                arrowprops=dict(arrowstyle="->", color=COLORS["red"], lw=1.5))
    ax.text(7.2, 1.75, "1:N", fontsize=6, color=COLORS["red"])

    path = os.path.join(IMG_DIR, "er_diagram.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def generate_request_flow_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(12, 5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Request Flow Diagram", fontsize=14, fontweight="bold", pad=20)

    steps = [
        (0.2, "Browser", COLORS["light_blue"], COLORS["blue"]),
        (2.2, "Nginx\n(TLS + RL)", COLORS["light_orange"], COLORS["orange"]),
        (4.2, "Gunicorn\n(WSGI)", COLORS["light_green"], COLORS["green"]),
        (6.2, "Django\n(Auth + View)", COLORS["light_purple"], COLORS["purple"]),
        (8.2, "ORM\n(Query)", COLORS["light_gray"], COLORS["gray"]),
        (10.2, "PostgreSQL", COLORS["light_gray"], COLORS["gray"]),
    ]

    for x, label, fc, ec in steps:
        draw_rounded_box(ax, x, 2.5, 1.6, 1, label, fc, ec, fontsize=8, fontweight="bold")

    for i in range(len(steps) - 1):
        x1 = steps[i][0] + 1.6
        x2 = steps[i + 1][0]
        draw_arrow(ax, x1, 3, x2, 3)

    labels = [
        (1.4, "HTTPS"),
        (3.4, "HTTP"),
        (5.4, "WSGI"),
        (7.4, "ORM"),
        (9.4, "SQL"),
    ]
    for x, label in labels:
        ax.text(x, 3.7, label, fontsize=7, ha="center", color=COLORS["gray"], fontstyle="italic")

    draw_rounded_box(ax, 6.2, 0.8, 1.6, 0.8, "Plotly Dash\n(iframe)", COLORS["light_purple"], COLORS["purple"], fontsize=7)
    draw_rounded_box(ax, 8.2, 0.8, 1.6, 0.8, "Redis\n(Cache)", COLORS["light_red"], COLORS["red"], fontsize=7)
    draw_arrow(ax, 7, 2.5, 7, 1.65)
    draw_arrow(ax, 9, 2.5, 9, 1.65)

    path = os.path.join(IMG_DIR, "request_flow.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex.replace("#", ""))
    shading.set(qn("w:val"), "clear")


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "#1E3A5F")

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "#F3F4F6")

    return table


def build_docx(arch_img, dfd_img, etl_img, er_img, req_img):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # --- TITLE PAGE ---
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("XYZ Platform")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Design Document")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Brown Brothers Harriman & Co.\nInternal Investment Banking Platform")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")
    doc.add_paragraph("")

    add_styled_table(doc, ["Field", "Value"], [
        ["Version", "1.0"],
        ["Date", "2026-06-24"],
        ["Stack", "Django 4.2 | PostgreSQL 16 | Redis 7 | Celery 5 | Airflow 2.9 | Plotly Dash"],
        ["Status", "Active Development"],
        ["Classification", "Confidential & Proprietary"],
    ])

    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Overview",
        "2. System Architecture",
        "3. Data Flow Diagram",
        "4. Django Applications",
        "5. Data Model (Entity Relationship)",
        "6. ETL Pipelines (Airflow)",
        "7. Interactive Dashboards (Plotly Dash)",
        "8. REST API",
        "9. Infrastructure & Deployment",
        "10. CI/CD Pipeline",
        "11. Security",
        "12. Observability",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = doc.styles["List Number"]

    doc.add_page_break()

    # --- 1. OVERVIEW ---
    doc.add_heading("1. Overview", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "The XYZ Platform is an internal investment banking application for XYZ Corp. "
        "that provides client and account management, portfolio dashboards with real-time analytics, "
        "risk metrics computation (VaR, Sharpe, drawdown), and ETL pipeline monitoring via Apache Airflow."
    )

    doc.add_heading("1.2 Users", level=2)
    add_styled_table(doc, ["Role", "Access Level"], [
        ["Superuser (admin)", "Full Django admin, all views, DAG triggering"],
        ["Staff (relationship managers)", "Client views, portfolio dashboard, analytics"],
        ["Read-only analysts", "Dashboards and risk metric views"],
    ])

    doc.add_heading("1.3 Technology Stack", level=2)
    add_styled_table(doc, ["Layer", "Technology"], [
        ["Web Framework", "Django 4.2.13"],
        ["API", "Django REST Framework 3.15"],
        ["Database", "PostgreSQL 16"],
        ["Cache / Broker", "Redis 7"],
        ["Task Queue", "Celery 5.4 + django-celery-beat"],
        ["Orchestration", "Apache Airflow 2.9.2 (LocalExecutor)"],
        ["Dashboards", "Plotly Dash 2.17 via django-plotly-dash"],
        ["WebSockets", "Django Channels 4.1 + channels-redis"],
        ["Static Files", "WhiteNoise 6.7"],
        ["Containerization", "Docker Compose"],
        ["Reverse Proxy", "Nginx (TLS, rate limiting)"],
        ["CI/CD", "GitHub Actions"],
        ["Error Tracking", "Sentry (production)"],
    ])

    doc.add_page_break()

    # --- 2. SYSTEM ARCHITECTURE ---
    doc.add_heading("2. System Architecture", level=1)
    doc.add_paragraph(
        "The platform follows a layered architecture with Django serving as the application layer, "
        "PostgreSQL as the primary data store, Redis for caching and message brokering, "
        "and Airflow for ETL pipeline orchestration."
    )
    doc.add_picture(arch_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2.1 Request Flow", level=2)
    doc.add_paragraph(
        "All requests are TLS-terminated at Nginx, rate-limited, and proxied to Gunicorn. "
        "Django handles authentication, routing, and ORM queries against PostgreSQL. "
        "Plotly Dash apps are served as iframes, and WebSocket connections upgrade through Django Channels."
    )
    doc.add_picture(req_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- 3. DATA FLOW DIAGRAM ---
    doc.add_heading("3. Data Flow Diagram", level=1)
    doc.add_paragraph(
        "Data flows from external sources (market data providers, custody systems) through Airflow ETL pipelines "
        "into PostgreSQL data stores, and is then consumed by the Django presentation layer (dashboards, views, APIs) "
        "for display to end users."
    )
    doc.add_picture(dfd_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- 4. DJANGO APPLICATIONS ---
    doc.add_heading("4. Django Applications", level=1)

    doc.add_heading("4.1 Accounts App", level=2)
    doc.add_paragraph(
        "Core entity hierarchy: Client -> Account -> Holding / Transaction. "
        "Provides list/detail views for clients and accounts, plus full CRUD via DRF API."
    )
    add_styled_table(doc, ["URL Pattern", "View", "Purpose"], [
        ["/clients/", "ClientListView", "List active clients, filter by name/tier"],
        ["/clients/<pk>/", "ClientDetailView", "Client profile + accounts + recent transactions"],
        ["/clients/account/<num>/", "AccountDetailView", "Holdings and transaction history"],
    ])

    doc.add_heading("4.2 Portfolio App", level=2)
    doc.add_paragraph(
        "Aggregated portfolio metrics with daily AUM snapshots. Main landing page embeds Plotly Dash."
    )
    add_styled_table(doc, ["URL Pattern", "View", "Purpose"], [
        ["/ and /dashboard/", "PortfolioDashboardView", "KPI cards + Plotly Dash iframe"],
        ["/api/snapshot/", "PortfolioSnapshotAPIView", "JSON AUM trend data (default 90 days)"],
    ])

    doc.add_heading("4.3 Analytics App", level=2)
    doc.add_paragraph(
        "Market data storage and risk metric computation. Provides dashboards for VaR, Sharpe, "
        "drawdown, and correlation analysis."
    )
    add_styled_table(doc, ["URL Pattern", "View", "Purpose"], [
        ["/analytics/", "AnalyticsDashboardView", "Risk KPI cards + Plotly Dash iframe"],
        ["/analytics/risk/", "RiskMetricListView", "Paginated risk metrics"],
        ["/analytics/market-data/<ticker>/", "MarketDataAPIView", "OHLCV JSON for a ticker"],
    ])

    doc.add_heading("4.4 ETL Monitor App", level=2)
    doc.add_paragraph(
        "Airflow pipeline monitoring with DAG run history, failure alerting, and manual trigger support. "
        "Celery tasks sync Airflow state every 5 minutes."
    )
    add_styled_table(doc, ["URL Pattern", "View", "Purpose"], [
        ["/etl/", "ETLDashboardView", "Latest DAG runs + unacknowledged alerts"],
        ["/etl/runs/", "DAGRunListView", "Paginated DAG run history"],
        ["/etl/trigger/<dag_id>/", "TriggerDAGView", "POST to trigger a DAG (permission required)"],
        ["/etl/api/alerts/", "PipelineAlertsAPIView", "JSON unacknowledged alerts"],
    ])

    doc.add_page_break()

    # --- 5. DATA MODEL ---
    doc.add_heading("5. Data Model", level=1)
    doc.add_heading("5.1 Entity Relationship Diagram", level=2)
    doc.add_picture(er_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5.2 Key Enumerations", level=2)
    add_styled_table(doc, ["Field", "Values"], [
        ["Client.tier", "UHNW, HNW, MA, INST"],
        ["Client.risk_profile", "CONSERVATIVE, MODERATE, AGGRESSIVE, VERY_AGGRESSIVE"],
        ["Account.account_type", "DISC, ADV, CUST, TRUST, RET"],
        ["Holding.asset_class", "EQ, FI, ALT, CASH, RE, COMM"],
        ["Transaction.type", "BUY, SELL, DIV, INT, FEE, TFI, TFO, DEP, WIT"],
        ["RiskMetric.scope", "ACCOUNT, PORTFOLIO, SECURITY"],
        ["DAGRun.state", "queued, running, success, failed, skipped"],
        ["PipelineAlert.severity", "CRITICAL, WARNING, INFO"],
    ])

    doc.add_heading("5.3 Database Indexes", level=2)
    add_styled_table(doc, ["Table", "Index", "Purpose"], [
        ["Client", "(tier, is_active)", "Filter clients by tier"],
        ["Holding", "(account, ticker) UNIQUE", "Prevent duplicate positions"],
        ["Transaction", "(account, trade_date)", "Query recent trades per account"],
        ["Transaction", "(ticker, trade_date)", "Cross-account ticker lookup"],
        ["MarketData", "(ticker, -price_date)", "Latest price fetch"],
        ["RiskMetric", "(scope, ref_id, date, lookback) UNIQUE", "Prevent duplicate metrics"],
        ["DAGRun", "(dag_id, -execution_date)", "Latest runs per DAG"],
    ])

    doc.add_page_break()

    # --- 6. ETL PIPELINES ---
    doc.add_heading("6. ETL Pipelines (Airflow)", level=1)
    doc.add_paragraph(
        "Three DAGs orchestrate the data pipeline. All use PostgreSQL connections and follow an "
        "extract-validate-transform-load pattern with idempotent upserts (ON CONFLICT DO UPDATE)."
    )

    doc.add_heading("6.1 Pipeline Flowcharts", level=2)
    doc.add_picture(etl_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("6.2 DAG Schedule Summary", level=2)
    add_styled_table(doc, ["DAG", "Schedule", "Timeout", "Retries"], [
        ["portfolio_etl_dag", "Mon-Fri 18:30 ET", "2 hours", "2 x 5 min"],
        ["market_data_dag", "Every 15 min (market hours)", "10 minutes", "3 x 2 min"],
        ["risk_report_dag", "Tue-Sat 00:00 UTC", "3 hours", "1 x 10 min"],
    ])

    doc.add_heading("6.3 Pipeline Dependency Chain", level=2)
    doc.add_paragraph(
        "Market Data DAG (intraday) -> Portfolio ETL DAG (18:30 ET) -> Risk Report DAG (19:00 ET). "
        "The Risk Report DAG uses an ExternalTaskSensor to wait for Portfolio ETL completion."
    )

    doc.add_heading("6.4 Stress Test Scenarios", level=2)
    add_styled_table(doc, ["Scenario", "Equity Shock", "Credit Spread", "Volatility"], [
        ["GFC 2008", "-42%", "+5%", "2.5x"],
        ["COVID 2020", "-34%", "+2.5%", "4x"],
        ["Taper Tantrum 2013", "-6%", "+130 bps rates", "-"],
        ["Rate Rise 300bps", "-18%", "+300 bps rates", "-"],
    ])

    doc.add_page_break()

    # --- 7. DASHBOARDS ---
    doc.add_heading("7. Interactive Dashboards (Plotly Dash)", level=1)

    add_styled_table(doc, ["Dashboard", "Mount Path", "Key Components"], [
        ["PortfolioDashApp", "/django_plotly_dash/app/PortfolioDashApp/",
         "AUM trend, asset allocation sunburst, attribution bar, rolling return"],
        ["RiskAnalyticsApp", "/django_plotly_dash/app/RiskAnalyticsApp/",
         "VaR fan chart, drawdown, correlation heatmap, efficient frontier"],
        ["ETLMonitorApp", "/django_plotly_dash/app/ETLMonitorApp/",
         "Gantt timeline, success rate bar, duration box plot"],
    ])

    doc.add_page_break()

    # --- 8. REST API ---
    doc.add_heading("8. REST API", level=1)
    doc.add_paragraph("Base URL: /api/v1/ | Authentication: Session + Basic | Pagination: 50 per page")

    add_styled_table(doc, ["Endpoint", "ViewSet", "Methods", "Features"], [
        ["/api/v1/accounts/clients/", "ClientViewSet", "CRUD", "Search, ordering"],
        ["/api/v1/accounts/clients/{id}/aum_summary/", "Custom action", "GET", "Returns total AUM"],
        ["/api/v1/accounts/accounts-list/", "AccountViewSet", "CRUD", "Filter, ordering"],
        ["/api/v1/accounts/transactions/", "TransactionViewSet", "Read-only", "Filter, ordering"],
        ["/api/v1/analytics/market-data/", "MarketDataViewSet", "Read-only", "Search, filter"],
        ["/api/v1/analytics/risk-metrics/", "RiskMetricViewSet", "Read-only", "Filter, ordering"],
    ])

    doc.add_page_break()

    # --- 9. INFRASTRUCTURE ---
    doc.add_heading("9. Infrastructure & Deployment", level=1)

    doc.add_heading("9.1 Docker Services", level=2)
    add_styled_table(doc, ["Service", "Image", "Port (Dev)", "Purpose"], [
        ["db", "postgres:16-alpine", "3000", "Primary database"],
        ["redis", "redis:7-alpine", "3001", "Cache + Celery broker"],
        ["django", "Custom (Python 3.11)", "3002", "Web application"],
        ["celery-worker", "Same as django", "-", "Background tasks"],
        ["celery-beat", "Same as django", "-", "Periodic scheduler"],
        ["airflow-webserver", "Custom (Airflow 2.9)", "3003", "Airflow UI"],
        ["airflow-scheduler", "Same as airflow", "-", "DAG execution"],
        ["nginx", "Custom (nginx)", "80, 443", "Reverse proxy (prod)"],
    ])

    doc.add_heading("9.2 Production Configuration", level=2)
    add_styled_table(doc, ["Component", "Setting"], [
        ["Gunicorn", "8 workers, gthread, 4 threads/worker, 120s timeout"],
        ["Celery", "8 concurrency"],
        ["PostgreSQL", "shared_buffers=512MB, effective_cache_size=2GB"],
        ["Redis", "Password-protected, persistence enabled"],
        ["Nginx", "TLS 1.2/1.3, rate limiting (5 login/min, 60 API/min)"],
        ["Django", "Non-root user (xyz:xyz), static collected at build"],
    ])

    doc.add_page_break()

    # --- 10. CI/CD ---
    doc.add_heading("10. CI/CD Pipeline", level=1)
    doc.add_paragraph("Platform: GitHub Actions (.github/workflows/ci-cd.yml)")

    add_styled_table(doc, ["Stage", "Trigger", "Actions"], [
        ["Lint", "All pushes/PRs", "Black, Flake8 (max-line=120), isort"],
        ["Test", "After lint", "pytest + coverage, PostgreSQL + Redis services"],
        ["Build", "Push only", "Multi-arch Docker build, push to GHCR"],
        ["Deploy Staging", "develop branch", "SSH pull + docker compose up + migrate"],
        ["Deploy Production", "main branch", "Blue-green deploy, Slack notification"],
    ])

    doc.add_page_break()

    # --- 11. SECURITY ---
    doc.add_heading("11. Security", level=1)

    doc.add_heading("11.1 Application Security", level=2)
    add_styled_table(doc, ["Control", "Implementation"], [
        ["Authentication", "Django session auth, @login_required on all views"],
        ["CSRF", "CsrfViewMiddleware + CSRF_TRUSTED_ORIGINS"],
        ["CORS", "django-cors-headers, whitelist-only in production"],
        ["Secrets", "Environment variables via .env, never committed"],
        ["Password Policy", "4 validators (similarity, min length, common, numeric)"],
    ])

    doc.add_heading("11.2 Production Hardening", level=2)
    add_styled_table(doc, ["Control", "Setting"], [
        ["HSTS", "Enabled with preload"],
        ["Secure Cookies", "SESSION_COOKIE_SECURE=True, CSRF_COOKIE_SECURE=True"],
        ["SSL Redirect", "SECURE_SSL_REDIRECT=True"],
        ["X-Frame-Options", "DENY (prod), SAMEORIGIN (dev for Dash iframes)"],
        ["Rate Limiting", "Login: 5 req/min, API: 60 req/min (Nginx)"],
        ["Container Security", "Non-root user (xyz:xyz), minimal base image"],
    ])

    doc.add_page_break()

    # --- 12. OBSERVABILITY ---
    doc.add_heading("12. Observability", level=1)

    doc.add_heading("12.1 Logging", level=2)
    add_styled_table(doc, ["Logger", "Level", "Handlers"], [
        ["Root", "INFO", "Console"],
        ["django", "WARNING", "Console + Rotating File"],
        ["apps", "DEBUG", "Console + Rotating File"],
        ["dags", "INFO", "Console + Rotating File"],
    ])
    doc.add_paragraph("Log files rotate at 10 MB with 5 backups (logs/xyz_platform.log).")

    doc.add_heading("12.2 Health Checks", level=2)
    add_styled_table(doc, ["Component", "Method"], [
        ["Django", "curl /admin/login/ (Docker healthcheck)"],
        ["Nginx", "/health/ -> 200 OK"],
        ["PostgreSQL", "pg_isready"],
        ["Redis", "redis-cli ping"],
        ["Airflow", "curl /health"],
    ])

    doc.add_heading("12.3 Error Tracking", level=2)
    doc.add_paragraph(
        "Sentry integration in production (if SENTRY_DSN is set). "
        "Celery task failures are logged and retried. "
        "Airflow failures auto-create CRITICAL PipelineAlert records."
    )

    # --- FOOTER ---
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "© Brown Brothers Harriman & Co. — Confidential & Proprietary. For internal use only."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.italic = True

    doc.save(OUTPUT_DOCX)
    print(f"Document saved to: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


def main():
    print("Generating diagrams...")
    arch_img = generate_system_architecture_diagram()
    print(f"  System Architecture: {arch_img}")

    dfd_img = generate_data_flow_diagram()
    print(f"  Data Flow Diagram: {dfd_img}")

    etl_img = generate_etl_flowchart()
    print(f"  ETL Flowchart: {etl_img}")

    er_img = generate_er_diagram()
    print(f"  ER Diagram: {er_img}")

    req_img = generate_request_flow_diagram()
    print(f"  Request Flow: {req_img}")

    print("\nBuilding Word document...")
    output = build_docx(arch_img, dfd_img, etl_img, er_img, req_img)
    print(f"\nDone! Output: {output}")


if __name__ == "__main__":
    main()
